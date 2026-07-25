import json
import zipfile

from docx import Document

from src.editorial.anonymization import (
    anonymize_text,
    create_anonymized_artifacts,
)


def test_anonymization_redacts_common_identity_fields() -> None:
    text = (
        "作者：张三\n"
        "单位：某某大学法学院\n"
        "邮箱 zhangsan@example.com，电话 13812345678。\n"
        + "正文讨论中国法治实践。"
        * 30
    )

    result = anonymize_text(text)

    assert "张三" not in result.text
    assert "某某大学" not in result.text
    assert "zhangsan@example.com" not in result.text
    assert "13812345678" not in result.text
    assert result.redaction_counts == {
        "email": 1,
        "phone": 1,
        "labeled_identity": 2,
    }


def test_short_anonymized_text_requires_editor_confirmation() -> None:
    result = anonymize_text("正文很短")

    assert result.requires_confirmation is True


def test_docx_anonymous_artifacts_remove_metadata_and_keep_table_and_footnote(
    tmp_path,
) -> None:
    source = tmp_path / "paper.docx"
    document = Document()
    document.core_properties.author = "张三"
    document.add_paragraph("论文标题", style="Title")
    document.add_paragraph("张三")
    document.add_paragraph("上海交通大学凯原法学院")
    document.add_paragraph("作者简介 张三，法学博士。")
    document.add_paragraph("正文讨论规范解释。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "材料"
    table.cell(0, 1).text = "结论"
    document.save(source)
    footnotes = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:footnote w:id="1"><w:p><w:r><w:t>脚注材料</w:t></w:r></w:p></w:footnote>
    </w:footnotes>"""
    with zipfile.ZipFile(source, "a") as archive:
        archive.writestr("word/footnotes.xml", footnotes)

    text_path, view_path, result, _, _ = create_anonymized_artifacts(
        str(source),
        "submission-1",
        root=tmp_path / "derived",
    )

    text = text_path.read_text(encoding="utf-8")
    view = json.loads(view_path.read_text(encoding="utf-8"))
    assert "张三" not in text
    assert "上海交通大学凯原法学院" not in text
    assert "作者简介" not in text
    assert "脚注材料" in text
    assert any(block["type"] == "table" for block in view["blocks"])
    assert any(block["type"] == "footnote" for block in view["blocks"])
    assert result.requires_confirmation is True
