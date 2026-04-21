from docx import Document
from src.ingestion.parsers.base import BaseParser, RawDocument
from src.core.exceptions import IngestionError


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> RawDocument:
        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return RawDocument(text=text, metadata={"format": "docx"})
        except Exception as e:
            raise IngestionError(f"DOCX 解析失败：{e}") from e
